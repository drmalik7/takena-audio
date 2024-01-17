#print('Hello world')
#name = input('What is your name?')
#print('Hello ' + name)

name = 'Chuck'
age = 17
pi = 3.14
cars = ['aston', 'bmw', 'mazda']

print(name)
print(age)
print(pi)
print(cars)

first_name = 'mike'
surname = 'murdoc'
full_name = first_name + ' ' + surname
print(full_name)
print(len(full_name))
full_name = first_name.capitalize() + ' ' + surname.capitalize()
print(full_name)

add = 10+5
subtract = 10-5
multiply = 10*5
divide = 10/5
mod = 10%5

print(add)
print(subtract)
print(multiply)
print(divide)
print(mod)

print(10<=10)
print(18>6)
print('Chuck'.endswith('d'))
print('Chuck'.endswith('k'))

isAdult=True
isTeenager=False
#age = 18

if isAdult:
    print('is Adult')
if age>=18:
    print('adult')
else:
    print('Teenager')
    
cars = ['aston', 'bmw', 'mazda', 'honda', 'tesla']
print(cars[0])
print(cars[3])

for car in cars:
    if car == 'bmw':
        print(car.capitalize())
    print(car)
   
   #class is a blue print 
class Person:
    def __init__(self, name, age): #defines two properties
        self.name = name 
        self.age = age
    
    #objects are variables in a class    
john=Person('John', 22)
mary = Person('Mary', 33)

#calling objects in a class
print(john.name+' '+str(john.age))
print(mary.name+' '+str(mary.age))

class Person:
    def __init__(self, name, age): #defines two properties
        self.name = name 
        self.age = age
    
    def walk(self):#defines behaviour of the properties
        print(self.name+' is walking...')
        
    def speak(self):
        print(' Hello my name is '+ self.name+ ' and i am '+ str(self.age)+ ' years old ')

   #objects are variables in a class    
john=Person('John', 22)
mary = Person('Mary', 33)
        
#calling objects in a class
print(john.name+' '+str(john.age))
john.speak()
john.walk()
print(mary.name+' '+str(mary.age))
mary.speak()
mary.walk()

#create documents
#install: pip3 install python-docx

from docx import Document

document = Document()

name = 'Chuck'
phone_number = 7999
email = 'chu@gmail.com'

document.add_paragraph(
    name + ' | ' + str(phone_number) + ' | ' + email + ' | ')

document.save('cv.docx')

'''from docx import Document

document = Document()

name = input('What is your name?')
phone_number = input('What is your phone number?')
email = input('What is your emaill address?')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email + ' | ')

document.save('cv.docx')'''
from docx import Document
#from docx import Inches
from pyttsx3 import Document

pyttsx3.speak('Hello')

document = Document()

#profile picture
document.add_picture('starguide.jpg')

name = input('What is your name? ')
phone_number = input('What is your phone number? ')
email = input('What is your email address? ')

#profile

document.add_paragraph(
    name + ' | ' + str(phone_number) + ' | ' + email + ' | ')

#about me

document.add_heading('About me')
about_me = input('Tell me about yourself ')
document.add_paragraph(about_me)

#work experience

document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('from Date ')
to_date = input('To Date ')

p.add_run(company + '' + '\n').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company)
p.add_run(experience_details)

# more experiences

while True:
    has_more_experiences = input(
        'Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company ')
        from_date = input('from Date ')
        to_date = input('To Date ')

        p.add_run(company + '' + '\n').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company)
        p.add_run(experience_details)
    else:
        break

#skills
document.add_heading('Skills')
skill = input('Enter Skills')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills? Yes or No ')
    if has_more_skills == 'yes':
        skill = input('Enter Skills')
        p = document.add_paragraph(skill)
        p.style ='List Bullet'
    else:
        break
    
#footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'Resume generated using hossanah and grace projects'



document.save('resume.docx')
