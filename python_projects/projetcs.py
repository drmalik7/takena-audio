'''from docx import Document

document = Document()

name = 'Andrew'
phone_number = 89457
email = 'andyfire@gmail.com'

document.add_paragraph(
    name + ' | ' + str(phone_number) + ' | ' + email + ' | ')

document.save('resume.docx')'''

from docx import Document
#from docx import Inches

document = Document()

#profile picture
document.add_picture('starguide.jpg')

name = input('What is your name? ')
phone_number = input('What is your phone number? ')
email = input('What is your email address? ')

#profile

document.add_paragraph(
    name + ' | ' + str(phone_number) + ' | ' + email + ' | ')

#about me

document.add_heading('About me')
about_me = input('Tell me about yourself ')
document.add_paragraph('about_me')

#work experience

document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company ')
from_date = input('from Date ')
to_date = input('To Date ')

p.add_run(company + ' ' + '\n').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company)
p.add_run('Experience Details')

document.save('rs.docx')